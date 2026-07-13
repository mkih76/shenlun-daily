#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
申论议论文 Word 文档生成器
公文字体规范：标题黑体三号，正文仿宋_GB2312三号，行距28磅
"""
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = Path(__file__).resolve().parent
ARTICLE_FILE = PROJECT_DIR / "today_articles.json"
CONFIG_FILE = PROJECT_DIR / "category_urls.json"
OUTPUT_DIR = PROJECT_DIR / "output"


# ---------- 字体工具 ----------
def set_run_font(run, font_name="仿宋_GB2312", size_pt=16, bold=False, color=None):
    """
    设置中英文字体（中文用 font_name，英文兼容）
    size_pt 默认 16 = 三号
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 中文显式设置（关键步骤，否则中文不应用）
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_paragraph_format(para, line_spacing_pt=28, first_line_indent_chars=2,
                         space_before=0, space_after=0, alignment=None):
    """段落格式：行距、首行缩进、对齐"""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent_chars:
        # 首行缩进 N 字符
        para.paragraph_format.first_line_indent = Pt(size_pt := 16 * first_line_indent_chars)
    if alignment is not None:
        para.alignment = alignment


def add_horizontal_line(doc):
    """添加分割线"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)


# ---------- 主体 ----------
def make_doc(articles, config):
    doc = Document()

    # 全局默认样式
    style = doc.styles["Normal"]
    style.font.name = "仿宋_GB2312"
    style.font.size = Pt(16)  # 三号
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    rFonts.set(qn("w:ascii"), "仿宋_GB2312")
    rFonts.set(qn("w:hAnsi"), "仿宋_GB2312")

    # 页面边距（公文标准）
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    # ===== 封面 =====
    today = datetime.now().strftime("%Y年%m月%d日")

    # 顶部留白
    spacer1 = doc.add_paragraph()
    run = spacer1.add_run("\n\n\n\n")
    set_run_font(run, "黑体", 14)
    spacer1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 装饰上横线
    deco1 = doc.add_paragraph()
    deco1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = deco1.add_run("━" * 20)
    set_run_font(run, "宋体", 18, color="C00000", bold=True)

    # 主标题
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("申论·议论文素材")
    set_run_font(run, "黑体", 44, bold=True, color="C00000")
    set_paragraph_format(title2, line_spacing_pt=60, first_line_indent_chars=0, space_after=8)

    # 装饰中横线
    deco2 = doc.add_paragraph()
    deco2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = deco2.add_run("—— 每日精选汇编 ——")
    set_run_font(run, "楷体_GB2312", 18, color="404040")
    set_paragraph_format(deco2, line_spacing_pt=24, first_line_indent_chars=0, space_after=8)

    # 装饰下横线
    deco3 = doc.add_paragraph()
    deco3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = deco3.add_run("━" * 20)
    set_run_font(run, "宋体", 18, color="C00000", bold=True)

    # 日期
    spacer2 = doc.add_paragraph()
    run = spacer2.add_run("\n\n")
    set_run_font(run, "黑体", 14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(today)
    set_run_font(run, "黑体", 32, bold=True, color="404040")
    set_paragraph_format(date_p, line_spacing_pt=42, first_line_indent_chars=0, space_after=12)

    weekday = ["星期一","星期二","星期三","星期四","星期五","星期六","星期日"][datetime.now().weekday()]
    weekday_p = doc.add_paragraph()
    weekday_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = weekday_p.add_run(weekday)
    set_run_font(run, "楷体_GB2312", 16, color="595959")
    set_paragraph_format(weekday_p, line_spacing_pt=24, first_line_indent_chars=0, space_after=24)

    # 统计
    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats_p.add_run(f"共收录 {len(articles)} 篇 议论文")
    set_run_font(run, "黑体", 16, bold=True, color="C00000")
    set_paragraph_format(stats_p, line_spacing_pt=24, first_line_indent_chars=0, space_after=6)

    stats_p2 = doc.add_paragraph()
    stats_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats_p2.add_run("来源：人民网·观点频道")
    set_run_font(run, "仿宋_GB2312", 14, color="595959")
    set_paragraph_format(stats_p2, line_spacing_pt=22, first_line_indent_chars=0, space_after=18)

    # 分类标签
    cat_p = doc.add_paragraph()
    cat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cat_p.paragraph_format.space_before = Pt(12)
    for i, cat in enumerate(articles.keys()):
        if i > 0:
            r = cat_p.add_run("   ")
            set_run_font(r, "楷体_GB2312", 12)
        cat_cfg = config["categories"][cat]
        r = cat_p.add_run(f"■ {cat_cfg['label']}")
        set_run_font(r, "黑体", 13, bold=True, color=cat_cfg["color"])
    set_paragraph_format(cat_p, line_spacing_pt=22, first_line_indent_chars=0, space_after=24)

    # 底部口号
    tip_p = doc.add_paragraph()
    tip_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tip_p.paragraph_format.space_before = Pt(36)
    run = tip_p.add_run("『 读时政评论 · 养申论语感 』")
    set_run_font(run, "楷体_GB2312", 14, color="808080", bold=False)
    set_paragraph_format(tip_p, line_spacing_pt=22, first_line_indent_chars=0)

    # 分页
    doc.add_page_break()

    # ===== 目录 =====
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    set_run_font(run, "黑体", 24, bold=True, color="C00000")
    set_paragraph_format(toc_title, line_spacing_pt=32, first_line_indent_chars=0, space_after=18)

    add_horizontal_line(doc)

    for i, (cat, art) in enumerate(articles.items(), 1):
        cat_cfg = config["categories"][cat]
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  【{cat_cfg['label']}】 ")
        set_run_font(run, "黑体", 14, bold=True, color=cat_cfg["color"])
        run2 = p.add_run(art["title"])
        set_run_font(run2, "仿宋_GB2312", 14, color="404040")
        set_paragraph_format(p, line_spacing_pt=24, first_line_indent_chars=0, space_after=4)

    doc.add_page_break()

    # ===== 正文 =====
    for i, (cat, art) in enumerate(articles.items(), 1):
        cat_cfg = config["categories"][cat]

        # 分类标签块
        badge = doc.add_paragraph()
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = badge.add_run(f"【 {cat_cfg['label']} · {cat_cfg['column_name']} 】")
        set_run_font(run, "黑体", 14, bold=True, color="FFFFFF")
        # 背景填充（用 shading）
        pPr = badge._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), cat_cfg["color"])
        pPr.append(shd)
        set_paragraph_format(badge, line_spacing_pt=24, first_line_indent_chars=0, space_after=12)

        # 标题
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(art["title"])
        set_run_font(run, "黑体", 22, bold=True, color="000000")
        set_paragraph_format(title_p, line_spacing_pt=36, first_line_indent_chars=0, space_before=6, space_after=10)

        # 元信息
        meta_parts = []
        if art.get("author"):
            meta_parts.append(f"作者：{art['author']}")
        if art.get("pub_date"):
            meta_parts.append(f"发表日期：{art['pub_date']}")
        meta_parts.append(f"栏目：{art.get('column_name', cat_cfg['column_name'])}")

        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_p.add_run("    ".join(meta_parts))
        set_run_font(run, "楷体_GB2312", 12, color="595959")
        set_paragraph_format(meta_p, line_spacing_pt=20, first_line_indent_chars=0, space_after=4)

        url_p = doc.add_paragraph()
        url_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = url_p.add_run(f"原文链接：{art['url']}")
        set_run_font(run, "楷体_GB2312", 10, color="808080")
        set_paragraph_format(url_p, line_spacing_pt=16, first_line_indent_chars=0, space_after=12)

        add_horizontal_line(doc)

        # 正文段落
        for para_text in art["content"].split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            # 清理单换行
            clean = re.sub(r"\n+", " ", para_text).strip()
            if len(clean) < 5:
                continue
            p = doc.add_paragraph()
            run = p.add_run(clean)
            set_run_font(run, "仿宋_GB2312", 16)
            set_paragraph_format(p, line_spacing_pt=28, first_line_indent_chars=2)

        # 篇末金句提示（如果有）
        if i < len(articles):
            doc.add_page_break()
        else:
            # 最后加收尾
            tail = doc.add_paragraph()
            tail.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = tail.add_run("\n\n—— 本日推送完 ——")
            set_run_font(run, "楷体_GB2312", 14, color="808080")
            set_paragraph_format(tail, line_spacing_pt=24, first_line_indent_chars=0, space_before=24)

    # ===== 页脚 =====
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"申论议论文素材 · {today} · 人民网观点频道 · 仅供学习使用")
    set_run_font(fr, "楷体_GB2312", 9, color="808080")

    return doc


def main():
    with open(ARTICLE_FILE, "r", encoding="utf-8") as f:
        articles = json.load(f)
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        config = json.load(f)

    if not articles:
        print("ERROR: today_articles.json is empty. Run crawl.py first.")
        return None

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    today = datetime.now().strftime("%Y-%m-%d")
    out_file = OUTPUT_DIR / f"{today}_每日议论文.docx"
    # 先写到临时文件，再原子重命名（避免 Word 占用导致权限错误）
    tmp_file = OUTPUT_DIR / f"{today}_每日议论文.tmp.docx"

    doc = make_doc(articles, config)
    doc.save(tmp_file)

    # 删除旧文件（如存在）+ 重命名
    if out_file.exists():
        try:
            out_file.unlink()
        except PermissionError:
            # Word 占用中：保留旧文件，新文件加时间戳后缀
            ts = datetime.now().strftime("%H%M%S")
            out_file = OUTPUT_DIR / f"{today}_每日议论文_{ts}.docx"
    tmp_file.rename(out_file)

    size_kb = out_file.stat().st_size / 1024
    print(f"✓ Generated: {out_file}")
    print(f"  Articles: {len(articles)}")
    print(f"  Size: {size_kb:.1f} KB")
    print(f"  Categories: {', '.join(articles.keys())}")
    return out_file


if __name__ == "__main__":
    main()